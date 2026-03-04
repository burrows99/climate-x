#!/usr/bin/env python3
"""
Extract contents from PDF, DOCX, and XLSX files and convert to Markdown
"""
import pymupdf  # PyMuPDF
from docx import Document
from openpyxl import load_workbook
import sys

def extract_pdf_to_md(pdf_path, output_md):
    """Extract text from PDF and write to markdown"""
    doc = pymupdf.open(pdf_path)
    with open(output_md, 'w', encoding='utf-8') as f:
        f.write(f"# {pdf_path.split('/')[-1]}\n\n")
        f.write("---\n\n")
        
        for page_num, page in enumerate(doc, start=1):
            f.write(f"## Page {page_num}\n\n")
            text = page.get_text()
            f.write(text)
            f.write("\n\n---\n\n")
    
    doc.close()
    print(f"✓ Extracted PDF to: {output_md}")

def extract_docx_to_md(docx_path, output_md):
    """Extract text from DOCX and write to markdown"""
    doc = Document(docx_path)
    with open(output_md, 'w', encoding='utf-8') as f:
        f.write(f"# {docx_path.split('/')[-1]}\n\n")
        f.write("---\n\n")
        
        for para in doc.paragraphs:
            if para.text.strip():
                # Try to preserve heading styles
                if para.style.name.startswith('Heading'):
                    level = para.style.name[-1] if para.style.name[-1].isdigit() else '1'
                    f.write(f"{'#' * int(level)} {para.text}\n\n")
                else:
                    f.write(f"{para.text}\n\n")
        
        # Extract tables
        if doc.tables:
            f.write("\n## Tables\n\n")
            for i, table in enumerate(doc.tables, start=1):
                f.write(f"### Table {i}\n\n")
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    f.write("| " + " | ".join(cells) + " |\n")
                    if table.rows[0] == row:
                        f.write("| " + " | ".join(["---"] * len(cells)) + " |\n")
                f.write("\n")
    
    print(f"✓ Extracted DOCX to: {output_md}")

def extract_xlsx_to_md(xlsx_path, output_md):
    """Extract data from XLSX and write to markdown"""
    wb = load_workbook(xlsx_path, data_only=True)
    with open(output_md, 'w', encoding='utf-8') as f:
        f.write(f"# {xlsx_path.split('/')[-1]}\n\n")
        f.write("---\n\n")
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            f.write(f"## Sheet: {sheet_name}\n\n")
            
            # Find the actual data range (non-empty cells)
            max_row = sheet.max_row
            max_col = sheet.max_column
            
            if max_row > 0 and max_col > 0:
                for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                    # Skip completely empty rows
                    if all(cell is None or str(cell).strip() == '' for cell in row):
                        continue
                    
                    cells = [str(cell) if cell is not None else "" for cell in row]
                    f.write("| " + " | ".join(cells) + " |\n")
                    
                    # Add separator after first row
                    if row_idx == 1:
                        f.write("| " + " | ".join(["---"] * len(cells)) + " |\n")
                
                f.write("\n")
            else:
                f.write("*Empty sheet*\n\n")
    
    print(f"✓ Extracted XLSX to: {output_md}")

if __name__ == "__main__":
    # Extract all files
    files = [
        ("Climate_X_-_Spectra_API_v2.0.pdf", "Climate_X_Spectra_API_v2.0.md", "pdf"),
        ("Solutions_Engineer_Challenge_July_2025_(3).docx", "Solutions_Engineer_Challenge.md", "docx"),
        ("CX_Task_Portfolio_Summary_(4).xlsx", "CX_Task_Portfolio_Summary.md", "xlsx")
    ]
    
    for input_file, output_file, file_type in files:
        try:
            if file_type == "pdf":
                extract_pdf_to_md(input_file, output_file)
            elif file_type == "docx":
                extract_docx_to_md(input_file, output_file)
            elif file_type == "xlsx":
                extract_xlsx_to_md(input_file, output_file)
        except Exception as e:
            print(f"✗ Error processing {input_file}: {e}")
    
    print("\n✓ All files processed!")
